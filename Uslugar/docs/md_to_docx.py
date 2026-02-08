# -*- coding: utf-8 -*-
"""Convert TEST-BLOCKS-MANIFEST-SPEC.md to Word .docx with tables, diagrams, and styling."""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Header row background for tables (light gray)
TABLE_HEADER_FILL = "E8E8E8"

def set_cell_shading(cell, fill):
    """Set table cell background color (hex without #)."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), fill)
    cell._tc.get_or_add_tcPr().append(shading)

def add_paragraph(doc, text, style='Normal'):
    p = doc.add_paragraph(style=style)
    if not text.strip():
        return p
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2] + ' ')
            run.bold = True
        else:
            p.add_run(part)
    return p

def is_table_row(line):
    line = line.strip()
    if '|' not in line:
        return False
    # Content between first and last pipe
    parts = [p.strip() for p in line.split('|')[1:-1]]
    return len(parts) > 0

def parse_table(lines, start):
    """Parse markdown table. Cells = content between pipes only (no empty edge columns)."""
    rows = []
    i = start
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if '|' not in stripped:
            break
        # Skip separator line (only dashes, pipes, spaces, colons)
        if re.match(r'^[\s|:\-]+$', stripped):
            i += 1
            continue
        # Cells = exactly what is between pipes: split('|')[1:-1], then strip
        cells = [c.strip() for c in line.split('|')[1:-1]]
        if not cells:
            i += 1
            continue
        rows.append(cells)
        i += 1
    return rows, i

def add_title_page(doc, lines):
    """Add title page from first lines; return index after ---."""
    doc.add_paragraph()
    title = "Specifikacija: Blokovski pristup testiranju"
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(18)
    doc.add_paragraph()
    meta = []
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if line.startswith('**Verzija:**'):
            meta.append(line.replace('**', ''))
        elif line.startswith('**Datum:**'):
            meta.append(line.replace('**', ''))
        elif line.startswith('**Status:**'):
            meta.append(line.replace('**', ''))
        elif line == '---':
            for m in meta:
                doc.add_paragraph(m)
            doc.add_paragraph()
            p = doc.add_paragraph()
            run = p.add_run()
            run.add_break(WD_BREAK.PAGE)
            return i + 1
        i += 1
    return 0

def insert_diagram(doc, img_path, caption=None, width_inches=5.0):
    """Insert image into document with optional caption."""
    if not img_path or not Path(img_path).exists():
        return
    try:
        doc.add_paragraph()
        doc.add_picture(str(img_path), width=Inches(width_inches))
        if caption:
            p = doc.add_paragraph(caption)
            p.paragraph_format.space_before = Pt(3)
            for r in p.runs:
                r.font.italic = True
                r.font.size = Pt(9)
                r.font.color.rgb = None  # default
        doc.add_paragraph()
    except Exception as e:
        doc.add_paragraph(f"[Slika: {img_path.name}]")

def main():
    base = Path(__file__).parent
    md_path = base / "TEST-BLOCKS-MANIFEST-SPEC.md"
    docx_path = base / "TEST-BLOCKS-MANIFEST-SPEC.docx"
    diagrams_dir = base / "doc_diagrams"

    # Generate diagrams
    try:
        from create_diagrams import create_all
        create_all()
    except Exception as e:
        print("Diagrams:", e)

    text = md_path.read_text(encoding='utf-8')
    lines = text.split('\n')

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    # Title page from first lines
    start_i = add_title_page(doc, lines)

    i = start_i
    in_code = False
    code_lines = []

    while i < len(lines):
        line = lines[i]
        line_stripped = line.strip()

        if line_stripped.startswith('```'):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                in_code = False
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph(code_text)
                for run in p.runs:
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                code_lines = []
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if line_stripped == '---':
            doc.add_paragraph('—' * 40)
            i += 1
            continue

        if line_stripped.startswith('### '):
            doc.add_heading(line_stripped[4:], level=3)
            doc.paragraphs[-1].paragraph_format.space_after = Pt(6)
            i += 1
            continue

        if line_stripped.startswith('## '):
            heading_text = line_stripped[3:]
            doc.add_heading(heading_text, level=2)
            doc.paragraphs[-1].paragraph_format.space_after = Pt(6)
            # Short detail callouts after selected sections
            if heading_text.startswith("1. ") and "Uvod" in heading_text:
                add_paragraph(doc, "Dokument pokriva: specifikaciju blokova i cigli, manifest (blocksManifest.js, blocksDefinitions.js), kontejnere kao test slučajeve, orkestrator (testRunnerService.js), API endpointe za testiranje i troubleshooting. Ciljna publika: razvojni i QA inženjeri, arhitekti.")
            elif heading_text.startswith("7. ") and "Datoteke" in heading_text:
                add_paragraph(doc, "Ključne datoteke: blocksManifest.js (BLOCKS_BY_TEST, getBlocksForTest), blocksDefinitions.js (BLOCK_DEFINITIONS), blocksValues.js (VALUES_BY_TEST), testRunnerService.js (runTestByBlocks, _runApiTest).")
            # Insert diagram after specific sections
            if heading_text.startswith("4. ") or ("Ulančavanje" in heading_text and "DAG" in heading_text):
                insert_diagram(doc, diagrams_dir / "dag_blocks.png", "Slika 1: Graf ovisnosti blokova (DAG) za test 4.1", width_inches=5.2)
            elif heading_text.startswith("8. ") or "Kontejner" in heading_text:
                insert_diagram(doc, diagrams_dir / "hierarchy.png", "Slika 2: Hijerarhija kontejner – blokovi – assert", width_inches=4.5)
            elif heading_text.startswith("15. ") or "Tok izvršavanja" in heading_text:
                insert_diagram(doc, diagrams_dir / "execution_flow.png", "Slika 3: Koraci izvršavanja testa (run-single)", width_inches=4.8)
            elif "Dodatak C" in heading_text or "Dijagrami toka" in heading_text:
                insert_diagram(doc, diagrams_dir / "registration_flow.png", "Slika 4: Tok testa 1.1 (registracija) – Client, Backend, Mailpit", width_inches=5.2)
            i += 1
            continue

        if line_stripped.startswith('# '):
            doc.add_heading(line_stripped[2:], level=1)
            doc.paragraphs[-1].paragraph_format.space_after = Pt(8)
            i += 1
            continue

        if line_stripped.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(line_stripped[2:].strip())
            i += 1
            continue

        prev_is_table = i > 0 and is_table_row(lines[i-1])
        if line_stripped and is_table_row(line) and not prev_is_table:
            rows, next_i = parse_table(lines, i)
            if rows:
                ncols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=ncols)
                table.style = 'Table Grid'
                for ri, row in enumerate(rows):
                    for ci in range(ncols):
                        cell = table.rows[ri].cells[ci]
                        cell_text = row[ci] if ci < len(row) else ""
                        cell.text = cell_text
                        if ri == 0:
                            for p in cell.paragraphs:
                                for r in p.runs:
                                    r.bold = True
                            set_cell_shading(cell, TABLE_HEADER_FILL)
                doc.add_paragraph()
                i = next_i
                continue

        if line_stripped:
            add_paragraph(doc, line_stripped)
        else:
            doc.add_paragraph()
        i += 1

    doc.save(docx_path)
    print(f"Saved: {docx_path}")

if __name__ == '__main__':
    main()
